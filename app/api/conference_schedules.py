# -*- coding: utf-8 -*-
"""Restful api for conference schedule."""

from flask import jsonify, request, url_for, Response
from flask_login import current_user
from sqlalchemy import or_
import requests
import random
import io
from docx import Document
from docx.shared import Pt
from .. import db
from ..models import Conference, ConferenceSchedule
from .errors import bad_request, forbidden
from .authentication import auth
from . import api


@api.route('/conference_schedules')
def get_conference_schedules():
    # for temperary use
    # r = requests.get('https://s3-ap-northeast-1.amazonaws.com/conferency-app/conferences.json')
    # return jsonify(r.json())
    schedule_list = {}
    for conference_schedule in ConferenceSchedule.query.filter_by(
            publish=True).all():
        conference = conference_schedule.conference
        schedule_list[conference.short_name] = {
            'title': conference.name,
            'short_title': conference.short_name.upper(),
            'about': conference.info,
            'location': conference.city + ', ' +
            (conference.state + ', ' if conference.state else '') +
            conference.country,
            'venue': conference.address,
            'time': {
                'start': str(conference.start_date) + ' 00:00:00',
                'end': str(conference.end_date) + ' 00:00:00',
            },
            'type': 'conference',
            'important': True,
            'data_url': url_for('api.get_conference_schedule_app',
                                short_name=conference.short_name,
                                _external=True),
            'timezone': conference.timezone
        }
        response = requests.get(
            'https://maps.googleapis.com/maps/api/geocode/json?address=' +
            conference.address + ',' +
            schedule_list[conference.short_name]['location'])
        resp_json_payload = response.json()
        if resp_json_payload['status'] == 'OK':
            schedule_list[conference.short_name]['latitude'] = resp_json_payload['results'][0]['geometry']['location']['lat']
            schedule_list[conference.short_name]['longitude'] = resp_json_payload['results'][0]['geometry']['location']['lng']
    return jsonify(schedule_list)


@api.route('/conference_schedules/app/<short_name>')
def get_conference_schedule_app(short_name):
    """Conference Schedule Api for mobile app."""
    # for temperary use
    conference = Conference.query.filter_by(short_name=short_name).first()
    if conference and conference.conference_schedule.publish:
        sessions_json = {}
        sessions = conference.conference_schedule.get_sessions.all()
        for session in sessions:
            sessions_json[session.id] = {
                'location': session.venue,
                'time': {
                    'start': str(session.start_time)[:16],
                    'end': str(session.end_time)[:16]
                },
                'title': session.title,
                'type': 'break' if session.type == 'regular' else session.type,
                'summary': session.description
            }
            if session.type != 'regular':
                speakers = session.speakers.all()
                if speakers:
                    sessions_json[session.id]['speakers'] = []
                    for speaker in speakers:
                        speaker_json = speaker.conference_profile_for_api(
                            conference)
                        sessions_json[session.id]['speakers'].append({
                            'avatar': 'https://app.conferency.com' +
                            speaker_json['avatar'] + '?' +
                            str(random.randint(0, 19)) if speaker_json['avatar'] else '',
                            'name': speaker_json['first_name'] + ' ' +
                            speaker_json['last_name'],
                            'organization': speaker_json['organization'],
                            'website': speaker_json['website'],
                            'summary': speaker_json['about_me'],
                            'email': speaker_json['email']
                        })
                moderators = session.moderators.all()
                if moderators:
                    sessions_json[session.id]['chairs'] = []
                    for moderator in moderators:
                        moderator_json = moderator.conference_profile_for_api(
                            conference)
                        sessions_json[session.id]['chairs'].append({
                            'avatar': 'https://app.conferency.com' +
                            moderator_json['avatar'] + '?' +
                            str(random.randint(0, 19)) if moderator_json['avatar'] else '',
                            'name': moderator_json['first_name'] + ' ' +
                            moderator_json['last_name'],
                            'organization': moderator_json['organization'],
                            'website': moderator_json['website'],
                            'summary': moderator_json['about_me'],
                            'email': moderator_json['email']
                        })
                if session.type == 'paper':
                    paper_sessions = session.paper_sessions.all()
                    if paper_sessions:
                        sessions_json[session.id]['papers'] = []
                        for paper_session in paper_sessions:
                            paper_json = {
                                'abstract': paper_session.paper.abstract,
                                'title': paper_session.paper.title,
                                'authors': [{
                                    'email': author.email,
                                    'name': author.full_name,
                                    'organization': author.organization,
                                    'website': author.website
                                } for author in paper_session.paper.authors_list]
                            }
                            discussants = paper_session.discussants.all()
                            if discussants:
                                paper_json['discussants'] = []
                                for discussant in discussants:
                                    discussant_json = discussant.conference_profile_for_api(conference)
                                    # print discussant_json
                                    paper_json['discussants'].append({
                                        'name':
                                        discussant_json['first_name'] + ' ' +
                                        discussant_json['last_name'],
                                        'email': discussant_json['email'],
                                        'organization':
                                        discussant_json['organization'],
                                        'website': discussant_json['website'],
                                        'summary': discussant_json['about_me'],
                                    })
                            sessions_json[session.id]['papers'].append(paper_json)
        return jsonify(sessions_json)
        # return jsonify(
        #     dict((session.id, session.to_json()) for session in sessions))
    else:
        return bad_request('Conference not found')


@api.route('/conference_schedules/<short_name>')
def get_conference_schedule(short_name):
    """Return conference schedule in json."""
    conference = Conference.query.filter_by(short_name=short_name).first()
    if conference:
        sessions = conference.conference_schedule.get_sessions.all()
        return jsonify(results=[session.to_json() for session in sessions])
    else:
        return bad_request('Conference not found')


@api.route('/conference_schedules/<short_name>/<type>')
def get_conference_schedule_file(short_name, type):
    """Download conference schedule."""
    conference = Conference.query.filter_by(short_name=short_name).first()
    if type not in ['txt', 'doc']:
        return bad_request('Invalid file type')
    if conference:
        sessions = conference.conference_schedule.get_sessions.all()
        if type == 'txt':
            txts = []
            for session in sessions:
                txts.append(
                    u'{session.start_time}\t{session.end_time}\t{session.title}\t\
                    {session.venue}'.format(session=session))
            return Response('\n'.join(txts), mimetype='text/plain',
                            headers={
                                'Content-disposition': 'attachment; \
                                filename={}_schedule.txt'.format(
                                    conference.short_name)})
        else:
            # only allow chair to download the doc
            if (current_user.is_anonymous) or \
                    (not current_user.is_chair(conference)):
                return forbidden('Not allowed')
            with io.BytesIO() as doc_stream:
                doc_document = Document()
                doc_document.add_heading(conference.name, 1)
                for session in sessions:
                    p = doc_document.add_paragraph(
                        u'{session.start_time} - {session.end_time}: '
                        '{session.title}'.format(
                            session=session),
                        style='ListBullet')
                    p.runs[0].font.size = Pt(13)
                    p.add_run('\nVenue: ' + session.venue)
                    speakers = session.speakers.all()
                    if speakers:
                        p.add_run('\nSpeakers:')
                        for speaker in speakers:
                            p.add_run(
                                '\n\t{speaker.full_name} '
                                '({speaker.organization})'.format(
                                    speaker=speaker))
                    moderators = session.moderators.all()
                    if moderators:
                        p.add_run('\nChairs:')
                        for moderator in moderators:
                            p.add_run(
                                '\n\t{moderator.full_name} '
                                '({moderator.organization})'.format(
                                    moderator=moderator))
                    papers = session.papers
                    if papers != []:
                        p.add_run('\nPapers:')
                        for paper in papers:
                            r = p.add_run('\n\t' + paper.title)
                            r.italic = True
                            r.bold = True
                            p.add_run('\n\t\t' + ', '.join(
                                [author.full_name for author in
                                    paper.authors_list]))
                doc_document.save(doc_stream)
                return Response(doc_stream.getvalue(),
                                mimetype='application/msword',
                                headers={
                                    'Content-disposition': 'attachment; \
                                    filename={}_schedule.doc'.format(
                                        conference.short_name)})
    else:
        return bad_request('Conference not found')


@api.route('/conference_schedules/<short_name>', methods=['PUT'])
@auth.login_required
def update_conference_schedule(short_name):
    """Update conference schedule config."""
    conference = Conference.query.filter_by(short_name=short_name).first()
    if conference:
        if not current_user.is_chair(conference):
            return forbidden('Not allowed')
        conference_schedule = conference.conference_schedule
        conference_schedule.publish = request.json.get('publish')
        db.session.add(conference_schedule)
        db.session.commit()
        return 'Success', 200
    else:
        return bad_request('Conference not found')


@api.route('/conference_schedules/app/search/<keyword>')
def search_conference_schedule(keyword):
    """Update conference schedule config."""
    conference_schedules = ConferenceSchedule.query.filter(
        Conference.id == ConferenceSchedule.conference_id,
        ConferenceSchedule.publish == True
    ).filter(
        or_(Conference.name.contains(keyword),
            Conference.short_name.contains(keyword),
            Conference.website.contains(keyword),
            Conference.contact_email.contains(keyword),
            Conference.address.contains(keyword),
            Conference.city.contains(keyword),
            Conference.state.contains(keyword),
            Conference.country.contains(keyword),
            Conference.info.contains(keyword),
            Conference.tags.contains(keyword),
            Conference.subjects.contains(keyword))
    ).all()
    schedule_list = {}
    for conference_schedule in conference_schedules:
        conference = conference_schedule.conference
        schedule_list[conference.short_name] = {
            'title': conference.name,
            'short_title': conference.short_name.upper(),
            'about': conference.info,
            'location': conference.city + ', ' +
            (conference.state + ', ' if conference.state else '') + conference.country,
            'venue': conference.address,
            'time': {
                'start': str(conference.start_date) + ' 00:00:00',
                'end': str(conference.end_date) + ' 00:00:00',
            },
            'type': 'conference',
            'important': True,
            'data_url': url_for('api.get_conference_schedule_app',
                                short_name=conference.short_name,
                                _external=True),
            'timezone': conference.timezone
        }
        response = requests.get('https://maps.googleapis.com/maps/api/geocode/json?address=' + conference.address + ',' + schedule_list[conference.short_name]['location'])
        resp_json_payload = response.json()
        if resp_json_payload['status'] == 'OK':
            schedule_list[conference.short_name]['latitude'] = resp_json_payload['results'][0]['geometry']['location']['lat']
            schedule_list[conference.short_name]['longitude'] = resp_json_payload['results'][0]['geometry']['location']['lng']
    return jsonify(schedule_list)
